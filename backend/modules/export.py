"""
GRACE Prototype — Document Export Helpers
Markdown → PDF (reportlab) and Markdown → DOCX (python-docx).
Pure-Python: no native libraries required.
"""
import io
import re
from html import escape

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Preformatted, Table, TableStyle, HRFlowable,
)
from reportlab.lib.enums import TA_LEFT

NAVY = colors.HexColor("#163265")
TEAL = colors.HexColor("#2A7A8A")
TEXT_DIM = colors.HexColor("#5A6F8C")
SOFT = colors.HexColor("#D5EDF2")


def _styles() -> dict:
    base = getSampleStyleSheet()
    return {
        "h1": ParagraphStyle("GraceH1", parent=base["Heading1"], textColor=NAVY,
            fontName="Helvetica-Bold", fontSize=18, spaceBefore=4, spaceAfter=10),
        "h2": ParagraphStyle("GraceH2", parent=base["Heading2"], textColor=TEAL,
            fontName="Helvetica-Bold", fontSize=14, spaceBefore=10, spaceAfter=6),
        "h3": ParagraphStyle("GraceH3", parent=base["Heading3"], textColor=NAVY,
            fontName="Helvetica-Bold", fontSize=12, spaceBefore=8, spaceAfter=4),
        "h4": ParagraphStyle("GraceH4", parent=base["Heading4"], textColor=TEAL,
            fontName="Helvetica-Bold", fontSize=11, spaceBefore=6, spaceAfter=3),
        "body": ParagraphStyle("GraceBody", parent=base["BodyText"], textColor=NAVY,
            fontName="Helvetica", fontSize=10.5, leading=14, alignment=TA_LEFT, spaceAfter=4),
        "bullet": ParagraphStyle("GraceBullet", parent=base["BodyText"], textColor=NAVY,
            fontName="Helvetica", fontSize=10.5, leading=14,
            leftIndent=14, bulletIndent=2, spaceAfter=2),
        "quote": ParagraphStyle("GraceQuote", parent=base["BodyText"], textColor=TEXT_DIM,
            fontName="Helvetica-Oblique", fontSize=10.5, leading=14,
            leftIndent=18, spaceBefore=4, spaceAfter=4),
        "code": ParagraphStyle("GraceCode", parent=base["Code"], textColor=NAVY,
            fontName="Courier", fontSize=9.5, leading=12,
            leftIndent=10, spaceAfter=6),
    }


def _inline_md_to_html(text: str) -> str:
    """Convert simple inline markdown to reportlab Paragraph HTML-like markup."""
    text = escape(text, quote=False)
    text = re.sub(r"\*\*([^*\n]+)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"<i>\1</i>", text)
    text = re.sub(r"`([^`\n]+)`", r'<font face="Courier">\1</font>', text)
    return text


def markdown_to_pdf_bytes(md_text: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title="GRACE Document",
    )
    s = _styles()
    flow = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            flow.append(Spacer(1, 4))
            i += 1
            continue

        if line.startswith("#### "):
            flow.append(Paragraph(_inline_md_to_html(line[5:]), s["h4"]))
        elif line.startswith("### "):
            flow.append(Paragraph(_inline_md_to_html(line[4:]), s["h3"]))
        elif line.startswith("## "):
            flow.append(Paragraph(_inline_md_to_html(line[3:]), s["h2"]))
        elif line.startswith("# "):
            flow.append(Paragraph(_inline_md_to_html(line[2:]), s["h1"]))
            flow.append(HRFlowable(width="100%", thickness=1, color=TEAL,
                                   spaceBefore=2, spaceAfter=8))
        elif re.match(r"^[\-\*]\s", line):
            flow.append(Paragraph("• " + _inline_md_to_html(line[2:].strip()), s["bullet"]))
        elif re.match(r"^\d+\.\s", line):
            m = re.match(r"^(\d+\.)\s+(.*)", line)
            if m:
                flow.append(Paragraph(f"{m.group(1)} {_inline_md_to_html(m.group(2))}", s["bullet"]))
        elif line.startswith("> "):
            flow.append(Paragraph(_inline_md_to_html(line[2:]), s["quote"]))
        elif line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            data = [[Paragraph(_inline_md_to_html(c), s["body"]) for c in header]] + [
                [Paragraph(_inline_md_to_html(c), s["body"]) for c in r] for r in rows
            ]
            tbl = Table(data, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), SOFT),
                ("TEXTCOLOR", (0, 0), (-1, 0), NAVY),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E5E7EB")),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            flow.append(tbl)
            flow.append(Spacer(1, 6))
            continue
        elif line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            flow.append(Preformatted("\n".join(code_lines), s["code"]))
        elif re.match(r"^[\-_*]{3,}$", line.strip()):
            flow.append(HRFlowable(width="100%", thickness=0.6,
                                   color=colors.HexColor("#E5E7EB"),
                                   spaceBefore=4, spaceAfter=4))
        else:
            flow.append(Paragraph(_inline_md_to_html(line), s["body"]))

        i += 1

    doc.build(flow)
    return buf.getvalue()


# ─── DOCX ────────────────────────────────────────────────────────────

_INLINE_RE = re.compile(r"(\*\*[^*\n]+\*\*|\*[^*\n]+\*|`[^`\n]+`)")


def _add_runs_with_inline_formatting(paragraph, text: str):
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) >= 3:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 3:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        else:
            paragraph.add_run(part)


def markdown_to_docx_bytes(md_text: str) -> bytes:
    doc = Document()
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        if line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif re.match(r"^[\-\*]\s", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_inline_formatting(p, line[2:].strip())
        elif re.match(r"^\d+\.\s", line):
            content = re.sub(r"^\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_inline_formatting(p, content)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(line[2:].strip())
            run.italic = True
        elif line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            ncols = len(header_cells)
            table = doc.add_table(rows=1 + len(rows), cols=ncols)
            table.style = "Light Grid"
            for c_idx, header in enumerate(header_cells):
                cell = table.rows[0].cells[c_idx]
                cell.text = ""
                run = cell.paragraphs[0].add_run(header)
                run.bold = True
            for r_idx, row in enumerate(rows):
                for c_idx in range(ncols):
                    cell_text = row[c_idx] if c_idx < len(row) else ""
                    table.rows[r_idx + 1].cells[c_idx].text = ""
                    _add_runs_with_inline_formatting(
                        table.rows[r_idx + 1].cells[c_idx].paragraphs[0], cell_text
                    )
            continue
        elif line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        elif re.match(r"^[\-_*]{3,}$", line.strip()):
            doc.add_paragraph("─" * 60)
        else:
            p = doc.add_paragraph()
            _add_runs_with_inline_formatting(p, line)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
